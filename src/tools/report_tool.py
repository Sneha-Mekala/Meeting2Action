# src/tools/report_tool.py
import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import mm

REPORTS_DIR = Path("artifacts/reports")
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

def _safe_name(meeting_id: str):
    ts = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    return f"{meeting_id}_{ts}"

def format_summary_lines(summary_text: str):
    # summary_text might be "- s1\n- s2..."
    lines = []
    for line in summary_text.splitlines():
        line = line.strip()
        if line.startswith("- "):
            lines.append(line[2:].strip())
        elif line:
            lines.append(line)
    return lines

class ReportTool:
    """
    Generates DOCX (Word), PDF and RTF meeting reports from structured pipeline output.
    """

    def generate(self, meeting_id: str, summary: str, actions: list, created_issues: list, notifications: list):
        """
        Returns dict with paths: {'docx':..., 'pdf':..., 'rtf':..., 'text': ...}
        """
        safe = _safe_name(meeting_id)
        docx_path = REPORTS_DIR / f"report_{safe}.docx"
        pdf_path = REPORTS_DIR / f"report_{safe}.pdf"
        rtf_path = REPORTS_DIR / f"report_{safe}.rtf"
        txt_path = REPORTS_DIR / f"report_{safe}.txt"

        # Build the "clean text" (plain but nicely formatted) for preview
        text = self._build_plain_text(meeting_id, summary, actions, created_issues, notifications)
        (REPORTS_DIR / txt_path.name).write_text(text, encoding="utf-8")

        # Create DOCX
        try:
            self._build_docx(docx_path, meeting_id, summary, actions, created_issues, notifications)
        except Exception as e:
            print("[ReportTool] DOCX generation error:", e)

        # Create PDF
        try:
            self._build_pdf(pdf_path, meeting_id, summary, actions, created_issues, notifications)
        except Exception as e:
            print("[ReportTool] PDF generation error:", e)

        # Create RTF (simple)
        try:
            self._build_rtf(rtf_path, meeting_id, summary, actions, created_issues, notifications)
        except Exception as e:
            print("[ReportTool] RTF generation error:", e)

        return {
            "report_text_path": str(txt_path),
            "report_docx_path": str(docx_path),
            "report_pdf_path": str(pdf_path),
            "report_rtf_path": str(rtf_path),
        }

    def _build_plain_text(self, meeting_id, summary, actions, created_issues, notifications):
        # Use the natural-language format you requested
        lines = []
        now = datetime.utcnow().strftime("%d %b %Y")
        lines.append(f"Meeting Summary – {meeting_id}")
        lines.append(f"Date: {now}")
        lines.append("")
        lines.append("Summary of Discussion:")
        lines.extend([f"- {s}" for s in format_summary_lines(summary)])
        lines.append("")
        lines.append("Extracted Action Items:")
        for i, a in enumerate(actions, start=1):
            owner = a.get("owner") or "Not assigned"
            due = a.get("due") or "Not specified"
            lines.append(f"{i}. {a.get('task')}")
            lines.append(f"   Owner: {owner}")
            lines.append(f"   Due: {due}")
            lines.append("")
        lines.append("Tasks Created:")
        for t in created_issues:
            assignee = t.get("assignee") or "No owner"
            due = t.get("due") or "Not specified"
            lines.append(f"- {t.get('id')} – {t.get('summary')} (Owner: {assignee}; Due: {due})")
        lines.append("")
        lines.append("Notifications Sent:")
        for n in notifications:
            lines.append(f"- Issue {n.get('issue')} -> status: {n.get('email_status')}")
        return "\n".join(lines)

    def _build_docx(self, path: Path, meeting_id, summary, actions, created_issues, notifications):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title
        title = doc.add_paragraph()
        run = title.add_run(f"Meeting Summary – {meeting_id}")
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x13, 0x47, 0x8A)  # professional blue
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph("")

        # Date / meta
        date_p = doc.add_paragraph()
        date_p.add_run("Date: ").bold = True
        date_p.add_run(datetime.utcnow().strftime("%d %b %Y"))
        doc.add_paragraph("")

        # Summary of Discussion
        h = doc.add_paragraph()
        run = h.add_run("Summary of Discussion")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x13, 0x47, 0x8A)
        doc.add_paragraph("")
        for s in format_summary_lines(summary):
            p = doc.add_paragraph(s, style='List Bullet')
            p_format = p.paragraph_format
            p_format.space_after = Pt(2)

        doc.add_paragraph("")

        # Action Items (as enumerated list with owner/due)
        h2 = doc.add_paragraph()
        run = h2.add_run("Extracted Action Items")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x13, 0x47, 0x8A)
        doc.add_paragraph("")
        for idx, a in enumerate(actions, start=1):
            p = doc.add_paragraph(f"{idx}. {a.get('task')}", style='List Number')
            # owner and due as indented lines
            p2 = doc.add_paragraph(f"Owner: {a.get('owner') or 'Not assigned'}")
            p2.paragraph_format.left_indent = Pt(18)
            p3 = doc.add_paragraph(f"Due: {a.get('due') or 'Not specified'}")
            p3.paragraph_format.left_indent = Pt(18)

        doc.add_paragraph("")

        # Tasks Created
        h3 = doc.add_paragraph()
        run = h3.add_run("Tasks Created")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x13, 0x47, 0x8A)
        doc.add_paragraph("")
        for t in created_issues:
            p = doc.add_paragraph(f"- {t.get('id')} — {t.get('summary')}")
            p2 = doc.add_paragraph(f"  Owner: {t.get('assignee') or 'No owner'}; Due: {t.get('due') or 'Not specified'}")
            p2.paragraph_format.left_indent = Pt(12)

        doc.add_paragraph("")

        # Notifications
        h4 = doc.add_paragraph()
        run = h4.add_run("Notifications Sent")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x13, 0x47, 0x8A)
        doc.add_paragraph("")
        for n in notifications:
            doc.add_paragraph(f"- Issue {n.get('issue')} -> status: {n.get('email_status')}")

        # Save
        doc.save(path)

    def _build_pdf(self, path: Path, meeting_id, summary, actions, created_issues, notifications):
        # Use a simple flowable document
        doc = SimpleDocTemplate(str(path), pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        styleN = styles['Normal']
        styleH = ParagraphStyle('Heading', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, textColor="#13478A")
        styleSub = ParagraphStyle('Sub', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, textColor="#13478A")
        styleBul = ParagraphStyle('Bul', parent=styleN, leftIndent=12, spaceAfter=6)

        story = []
        story.append(Paragraph(f"Meeting Summary – {meeting_id}", styleH))
        story.append(Spacer(1, 6))

        story.append(Paragraph(f"<b>Date:</b> {datetime.utcnow().strftime('%d %b %Y')}", styleN))
        story.append(Spacer(1, 8))

        story.append(Paragraph("Summary of Discussion", styleSub))
        for s in format_summary_lines(summary):
            story.append(Paragraph(f"• {s}", styleBul))
        story.append(Spacer(1, 8))

        story.append(Paragraph("Extracted Action Items", styleSub))
        for idx, a in enumerate(actions, start=1):
            story.append(Paragraph(f"{idx}. {a.get('task')}", styleN))
            story.append(Paragraph(f"Owner: {a.get('owner') or 'Not assigned'}; Due: {a.get('due') or 'Not specified'}", styleBul))
        story.append(Spacer(1, 8))

        story.append(Paragraph("Tasks Created", styleSub))
        for t in created_issues:
            story.append(Paragraph(f"• {t.get('id')} — {t.get('summary')} (Owner: {t.get('assignee') or 'No owner'}; Due: {t.get('due') or 'Not specified'})", styleBul))

        story.append(Spacer(1, 8))
        story.append(Paragraph("Notifications Sent", styleSub))
        for n in notifications:
            story.append(Paragraph(f"• Issue {n.get('issue')} -> status: {n.get('email_status')}", styleBul))

        doc.build(story)

    def _build_rtf(self, path: Path, meeting_id, summary, actions, created_issues, notifications):
        # Simple RTF writer with bold headings and blue color for headings
        def rtf_escape(s: str):
            return s.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
        blue_rgb = r'\red19\green71\blue138'
        parts = []
        parts.append(r'{\rtf1\ansi')
        parts.append(r'{\colortbl ;' + blue_rgb + ';}')
        parts.append(r'\fs24')  # default font size
        parts.append(r'\b\cf1 ' + rtf_escape(f"Meeting Summary – {meeting_id}") + r'\b0\par')
        parts.append(rtf_escape("Date: " + datetime.utcnow().strftime("%d %b %Y")) + r'\par\par')
        parts.append(r'\b\cf0 Summary of Discussion\b0\par')
        for s in format_summary_lines(summary):
            parts.append(rtf_escape("• " + s) + r'\par')
        parts.append(r'\par')
        parts.append(r'\b Extracted Action Items\b0\par')
        for idx, a in enumerate(actions, start=1):
            parts.append(rtf_escape(f"{idx}. {a.get('task')}") + r'\par')
            parts.append(rtf_escape(f"   Owner: {a.get('owner') or 'Not assigned'}") + r'\par')
            parts.append(rtf_escape(f"   Due: {a.get('due') or 'Not specified'}") + r'\par')
        parts.append(r'\par')
        parts.append(r'\b Tasks Created\b0\par')
        for t in created_issues:
            parts.append(rtf_escape(f"- {t.get('id')} — {t.get('summary')} (Owner: {t.get('assignee') or 'No owner'}; Due: {t.get('due') or 'Not specified'})") + r'\par')
        parts.append(r'\par')
        parts.append(r'\b Notifications Sent\b0\par')
        for n in notifications:
            parts.append(rtf_escape(f"- Issue {n.get('issue')} -> status: {n.get('email_status')}") + r'\par')
        parts.append('}')
        path.write_text("\n".join(parts), encoding="utf-8")
